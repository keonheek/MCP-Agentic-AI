"""
DOCX exporter for grant application drafts.
Requires: python-docx
"""
from pathlib import Path
from datetime import date


def export_to_docx(program: str, business_name: str, draft: dict, output_path: str = None) -> str:
    """
    Export draft dict to DOCX file.
    Falls back to plain text if python-docx not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f"{program} 지원사업 신청서", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"사업명: {business_name}")
        doc.add_paragraph(f"작성일: {date.today().isoformat()}")
        doc.add_paragraph(f"작성: 1stmover AI 초안 서비스")
        doc.add_paragraph("")

        for question, answer in draft.items():
            doc.add_heading(question, level=2)
            p = doc.add_paragraph(answer)
            p.style.font.size = Pt(11)
            doc.add_paragraph("")

        if not output_path:
            output_path = f"{business_name}_{program}_신청서_{date.today().isoformat()}.docx"

        doc.save(output_path)
        return output_path

    except ImportError:
        # Fallback: plain text
        if not output_path:
            output_path = f"{business_name}_{program}_신청서_{date.today().isoformat()}.txt"
        lines = [f"# {program} 신청서 - {business_name}\n", f"작성일: {date.today().isoformat()}\n\n"]
        for question, answer in draft.items():
            lines.append(f"## {question}\n{answer}\n\n")
        Path(output_path).write_text("".join(lines), encoding="utf-8")
        return output_path
