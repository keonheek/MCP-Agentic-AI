"""Build CV_KeonheeKim_Updated.docx — one page, no bold prefixes, hyperlinks on URLs"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "CV_KeonheeKim_Updated.docx")

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Inches(0.60)
    sec.bottom_margin = Inches(0.60)
    sec.left_margin   = Inches(0.70)
    sec.right_margin  = Inches(0.70)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)

RIGHT_TAB = int((8.5 - 1.4) * 72 * 20)
BODY = 10

def sp(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def set_font(run, size=BODY, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic

def add_right_tab(para):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(RIGHT_TAB))
    tabs.append(tab)
    pPr.append(tabs)

def _looks_like_url(s):
    s = s.strip()
    return (s.startswith("github.com") or s.startswith("http") or
            ".streamlit.app" in s or ".io" in s)

def _url_for(s):
    s = s.strip()
    return s if s.startswith("http") else "https://" + s

def add_hyperlink_run(para, text, url, size=BODY, italic=False):
    """Append a clickable hyperlink run to an existing paragraph."""
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if italic:
        rPr.append(OxmlElement("w:i"))
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text.strip()
    r.append(t)
    hl.append(r)
    para._p.append(hl)

def section_header(title):
    p = doc.add_paragraph()
    sp(p, before=4, after=1)
    r = p.add_run(title)
    set_font(r, bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)

def role_line(org, detail, date):
    p = doc.add_paragraph()
    sp(p, before=3, after=0)
    add_right_tab(p)
    r = p.add_run(org)
    set_font(r, bold=True)
    if detail:
        # Split on "  |  " so URLs can be made into hyperlinks
        parts = detail.split("  |  ")
        for i, part in enumerate(parts):
            sep = " | " if i == 0 else "  |  "
            if _looks_like_url(part):
                r_sep = p.add_run(sep)
                set_font(r_sep, italic=True)
                add_hyperlink_run(p, part.strip(), _url_for(part.strip()), italic=True)
            else:
                r2 = p.add_run(sep + part.strip())
                set_font(r2, italic=True)
    r3 = p.add_run("\t" + date)
    set_font(r3)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    sp(p, before=0, after=0)
    p.paragraph_format.left_indent       = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    rt = p.add_run(text)
    set_font(rt)

def skill_line(label, content):
    p = doc.add_paragraph()
    sp(p, before=0, after=0)
    p.paragraph_format.left_indent = Inches(0.12)
    rb = p.add_run(label)
    set_font(rb, bold=True)
    rt = p.add_run(content)
    set_font(rt)

# ── HEADER ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
sp(p, before=0, after=1)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Keonhee Kim")
set_font(r, size=15, bold=True)

p2 = doc.add_paragraph()
sp(p2, before=0, after=3)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    "Seoul, South Korea  |  (+82) 010-5704-4832  |  keonhee3337@gmail.com  |  "
)
set_font(r2)
add_hyperlink_run(p2, "github.com/keonhee3337-art", "https://github.com/keonhee3337-art")
r3 = p2.add_run("  |  Availability: Feb 24 \u2013 Sep 10")
set_font(r3)

# ── EDUCATION ──────────────────────────────────────────────────────────
section_header("EDUCATION")
role_line("Sungkyunkwan University (SKKU)", "Seoul, South Korea", "")

p = doc.add_paragraph()
sp(p, before=0, after=0)
set_font(p.add_run("Bachelor of Business Administration  |  "))
ri = p.add_run("Mar 2020 \u2013 Present (Expected Feb 2027)")
set_font(ri, italic=True)

bullet("GPA: 3.7 / 4.5  \u2022  Advanced Training: Financial Modeling, K-GAAP, Cost Accounting (Jan 2024 \u2013 Jun 2025)")

# ── PROJECT & TECHNICAL EXPERIENCE ─────────────────────────────────────
section_header("PROJECT & TECHNICAL EXPERIENCE")

# M&A Due Diligence Suite
role_line(
    "M\u0026A Due Diligence Suite",
    "AI Engineer  |  github.com/keonhee3337-art/consulting-emulation",
    "Jan 2026 \u2013 Present"
)
bullet(
    "Built a LangGraph Supervisor pipeline routing queries to four specialized agents \u2014 Text2SQL, RAG, "
    "DCF Valuation, and XGBoost Financial Distress \u2014 automating M\u0026A due diligence workflows modeled on MBB/Big4 practices."
)
bullet(
    "Integrated live Korean DART financial disclosure API (Samsung Electronics, SK Hynix, LG Electronics); "
    "delivers DCF + EV/EBITDA comps analysis in under 60 seconds per company."
)
bullet(
    "Deployed production REST API on AWS Lambda (ap-northeast-2) with Streamlit Cloud dashboard; "
    "automated PowerPoint deck and multi-sheet Excel DCF model export per query."
)

# FinAgent
role_line(
    "FinAgent",
    "AI Engineer  |  github.com/keonhee3337-art/FinAgent  |  keonhee-finagent.streamlit.app",
    "Feb 2026 \u2013 Present"
)
bullet(
    "Architected a 3-node LangGraph pipeline (RAG \u2192 Text2SQL \u2192 synthesis) reducing Samsung Electronics "
    "and SK Hynix research time from several hours to under 2 minutes."
)
bullet(
    "Built a custom VectorDB from scratch (OpenAI embeddings + NumPy cosine similarity) due to Python 3.14 / "
    "ChromaDB incompatibility; wrapped with LangChain BaseRetriever for full framework interoperability."
)

# RAG Demo
role_line(
    "RAG Demo",
    "Full-Stack Developer",
    "Feb 2026 \u2013 Mar 2026"
)
bullet(
    "Engineered a production RAG API (FastAPI + Pinecone + GPT-4o) with Supabase-backed conversation history; "
    "exposed live endpoint via ngrok for real-time client demonstrations."
)
bullet(
    "Implemented OpenAI text-embedding-ada-002 for document ingestion and semantic retrieval; "
    "integrated multi-turn session management via Supabase PostgreSQL."
)

# SKKU-Deloitte SDC
role_line(
    "SKKU-Deloitte AI/IT Consulting Society (SDC)",
    "Founder \u0026 President",
    "Sep 2025 \u2013 Present"
)
bullet(
    "Structured a talent partnership with Deloitte Partners to secure a formal MOU and internship pipeline; "
    "designed a competency-based L\u0026D curriculum accelerating analyst readiness by 40%."
)
bullet(
    "Built an automated applicant screening pipeline (Gmail \u2192 PDF extraction \u2192 Claude Haiku AI scoring \u2192 "
    "Google Sheets) reducing 50+ application reviews by 85% (~4 hours \u2192 30 minutes)."
)

# Samsung Electronics AI Strategy Engine
role_line(
    "Samsung Electronics AI Strategy Engine",
    "Full-Stack Developer  |  Live app available",
    "Jan 2025 \u2013 Feb 2025"
)
bullet(
    "Engineered an ETL pipeline (Python, DART-FSS, yfinance) extracting 7 years of audited financials into "
    "SQLite; built a RAG system using GPT-4o for real-time strategic database querying."
)
bullet(
    "Developed a Facebook Prophet ML model forecasting stock trends with 95% confidence intervals and an "
    "interactive macroeconomic shock simulator with Plotly waterfall charts; deployed on Streamlit Cloud."
)

# ── ADDITIONAL EXPERIENCE ───────────────────────────────────────────────
section_header("ADDITIONAL EXPERIENCE")

role_line("Ensight English Academy", "Seoul, South Korea  |  Academic Instructor \u0026 Speech Coach", "Sep 2023 \u2013 Mar 2024")
bullet(
    "Directed a rigorous communication curriculum mirroring executive standards, leading students to win the "
    "Grand Prize (Daesang) at the National English Speech Contest."
)

role_line("Independent Communications Coach", "Freelance", "Jun 2019 \u2013 Sep 2019")
bullet(
    "Diagnosed speech patterns of non-native speakers and executed 1:1 logic and delivery training, "
    "resulting in a 2nd Place win at the UC Berkeley Speech Contest."
)

# ── SKILLS & ADDITIONAL INFORMATION ────────────────────────────────────
section_header("SKILLS & ADDITIONAL INFORMATION")

skill_line("AI & Engineering: ",
    "LangGraph, multi-agent orchestration, RAG, custom VectorDB, Text2SQL, MCP servers, "
    "AWS Lambda, LangChain, OpenAI API, Claude API, XGBoost, Claude Code (11 skills, 7 sub-agents, hook automation)."
)
skill_line("Technical & Data: ",
    "Python (Pandas, Streamlit, FastAPI), SQL, Financial Modeling (Excel), Facebook Prophet, Plotly, GitHub CI/CD."
)
skill_line("Languages: ",
    "English (Native, 16 years overseas, OPIc: AL), Korean (Native)."
)
skill_line("Interests: ",
    "Classical Music (Concertmaster for European orchestra), Sports (Varsity Basketball Captain)."
)
skill_line("Military Service: ",
    "Republic of Korea Army (Sergeant); Completed Full Term of Service (2023)."
)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
